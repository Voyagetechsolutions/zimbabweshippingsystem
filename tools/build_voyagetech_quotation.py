from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Mthokozisi.DESKTOP-DPOBCC1\Documents\zimbabwe-shipping-nexus")
OUT = ROOT / "deliverables" / "Voyage_Tech_Zimbabwe_Shipping_Quotation.docx"
LOGO = Path(r"C:\Users\Mthokozisi.DESKTOP-DPOBCC1\Pictures\Voyagetech.png")

# standard_business_brief preset, with a named Voyage Tech brand-color override.
BLUE = "125A98"
DARK_BLUE = "0A3154"
PALE_BLUE = "EAF3FA"
MID_BLUE = "D5E8F6"
INK = "172534"
MUTED = "607080"
LIGHT = "F5F8FB"
WHITE = "FFFFFF"
BORDER = "CCD8E3"
GREEN = "0D7A51"
AMBER = "9A6500"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_run(run, size=10.5, color=INK, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = "left" if edge == "start" else "right" if edge == "end" else edge
        node = borders.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            borders.append(node)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                node.set(qn(f"w:{key}"), str(edge_data[key]))


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run(run, size=8.5, color=MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    keep_with_next(p)
    return p


def add_compact_heading(doc, text):
    """Named page-four override for a compact commercial-terms page."""
    p = doc.add_paragraph(text, style="Heading 1")
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(4)
    keep_with_next(p)
    return p


def add_body(doc, text, bold_lead=None, italic=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True)
        set_run(p.add_run(text[len(bold_lead):]), italic=italic)
    else:
        set_run(p.add_run(text), italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    set_run(p.add_run(text), size=10.2)
    return p


def add_feature_group(doc, heading, items):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    keep_with_next(p)
    set_run(p.add_run(heading), size=11, color=BLUE, bold=True)
    for item in items:
        add_bullet(doc, item)


def add_label_value(cell, label, value, value_color=INK, value_size=10.5):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    set_run(p.add_run(label + "\n"), size=8.5, color=MUTED, bold=True)
    set_run(p.add_run(value), size=value_size, color=value_color, bold=True)


def add_pricing_table(doc, rows, total_label, total_value, widths=(6900, 2460)):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, list(widths))
    hdr = table.rows[0].cells
    hdr[0].text = "DELIVERABLE"
    hdr[1].text = "PRICE (USD)"
    for cell in hdr:
        set_cell_shading(cell, DARK_BLUE)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=9, color=WHITE, bold=True)
    hdr[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_repeat_header(table.rows[0])
    border_spec = {"val": "single", "sz": "6", "space": "0", "color": BORDER}
    for idx, (label, value) in enumerate(rows):
        cells = table.add_row().cells
        if idx % 2 == 1:
            set_cell_shading(cells[0], LIGHT)
            set_cell_shading(cells[1], LIGHT)
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(label), size=10.2)
        p = cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(value), size=10.2, bold=True)
        for cell in cells:
            set_cell_border(cell, top=border_spec, bottom=border_spec, start=border_spec, end=border_spec)
    cells = table.add_row().cells
    for cell in cells:
        set_cell_shading(cell, PALE_BLUE)
        set_cell_border(cell, top={"val": "single", "sz": "14", "space": "0", "color": BLUE}, bottom=border_spec, start=border_spec, end=border_spec)
    p = cells[0].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(total_label), size=11.5, color=DARK_BLUE, bold=True)
    p = cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(total_value), size=12, color=BLUE, bold=True)
    return table


def add_terms_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [2500, 6860])
    border_spec = {"val": "single", "sz": "5", "space": "0", "color": BORDER}
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], PALE_BLUE)
        set_run(cells[0].paragraphs[0].add_run(label), size=9.5, color=DARK_BLUE, bold=True)
        set_run(cells[1].paragraphs[0].add_run(value), size=9.5)
        for cell in cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            set_cell_border(cell, top=border_spec, bottom=border_spec, start=border_spec, end=border_spec)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.font.color.rgb = rgb(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, before, after, color in (
    ("Heading 1", 16, 16, 8, BLUE),
    ("Heading 2", 13, 12, 6, BLUE),
    ("Heading 3", 12, 8, 4, DARK_BLUE),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = rgb(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

bullet = styles["List Bullet"]
bullet.font.name = "Calibri"
bullet.font.size = Pt(10.2)
bullet.paragraph_format.left_indent = Inches(0.5)
bullet.paragraph_format.first_line_indent = Inches(-0.25)
bullet.paragraph_format.space_after = Pt(4)
bullet.paragraph_format.line_spacing = 1.10

# Running header and footer.
header_p = section.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header_p.paragraph_format.space_after = Pt(0)
set_run(header_p.add_run("VOYAGE TECH  |  SOFTWARE QUOTATION"), size=8.5, color=MUTED, bold=True)

footer_table = section.footer.add_table(rows=1, cols=2, width=Inches(6.5))
set_table_geometry(footer_table, [7200, 2160], indent=0)
footer_table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
footer_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(footer_table.rows[0].cells[0].paragraphs[0].add_run("Voyage Tech | Confidential commercial quotation"), size=8.5, color=MUTED)
p = footer_table.rows[0].cells[1].paragraphs[0]
set_run(p.add_run("Page "), size=8.5, color=MUTED)
add_page_field(p)

# PAGE 1 - customer pack opening.
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
logo_run = p.add_run()
inline = logo_run.add_picture(str(LOGO), width=Inches(2.75))
inline._inline.docPr.set("descr", "Voyage Tech logo")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
set_run(p.add_run("COMMERCIAL QUOTATION"), size=10, color=BLUE, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(7)
set_run(p.add_run("Zimbabwe Shipping Mobile Apps"), size=27, color=DARK_BLUE, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(20)
set_run(p.add_run("Customer app + staff operations app"), size=14, color=MUTED)

meta = doc.add_table(rows=2, cols=2)
set_table_geometry(meta, [4680, 4680])
add_label_value(meta.cell(0, 0), "PREPARED FOR", "Zimbabwe Shipping")
add_label_value(meta.cell(0, 1), "PREPARED BY", "Voyage Tech")
add_label_value(meta.cell(1, 0), "QUOTATION NUMBER", "VT-ZS-2026-0716")
add_label_value(meta.cell(1, 1), "DATE / VALIDITY", "16 July 2026 / 30 days")
for row in meta.rows:
    for cell in row.cells:
        set_cell_shading(cell, LIGHT)
        set_cell_border(cell, top={"val": "single", "sz": "5", "color": BORDER}, bottom={"val": "single", "sz": "5", "color": BORDER}, start={"val": "single", "sz": "5", "color": BORDER}, end={"val": "single", "sz": "5", "color": BORDER})

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_heading(doc, "Project overview", 1)
add_body(doc, "Voyage Tech proposes the design, development, integration, testing and release preparation of two connected mobile applications for Zimbabwe Shipping. The customer app will make booking and tracking simple, while the staff app will give administrators, finance teams and drivers the tools needed to manage daily operations from mobile devices.")

callout = doc.add_table(rows=1, cols=2)
set_table_geometry(callout, [6100, 3260])
set_cell_shading(callout.cell(0, 0), PALE_BLUE)
set_cell_shading(callout.cell(0, 1), BLUE)
add_label_value(callout.cell(0, 0), "COMPLETE TWO-APP IMPLEMENTATION", "Fixed project investment", value_color=DARK_BLUE, value_size=12)
p = callout.cell(0, 1).paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(0)
set_run(p.add_run("USD\n"), size=9, color=WHITE, bold=True)
set_run(p.add_run("$1,199.99"), size=20, color=WHITE, bold=True)

doc.add_paragraph().paragraph_format.space_after = Pt(1)
add_body(doc, "Recommended ongoing service: $149.99 per month after launch. This covers hosting oversight, monitoring, backups, maintenance and a monthly allowance for minor updates.", italic=True, after=4)

# PAGE 2 - customer app.
doc.add_page_break()
add_heading(doc, "1. Customer mobile app", 1)
add_body(doc, "A simple, customer-focused application for people shipping goods between the United Kingdom and Zimbabwe.")

add_feature_group(doc, "Account and customer experience", [
    "Secure registration, sign-in, profile management and password recovery.",
    "Personal dashboard with greeting, recent activity and quick actions.",
    "Saved sender, receiver, collection and delivery details.",
])
add_feature_group(doc, "Shipment booking and quotations", [
    "Step-by-step shipment booking for collection and delivery requirements.",
    "Route, package, service and address capture with input validation.",
    "Instant standard pricing and custom quotation requests where required.",
    "Booking confirmation, reference number and booking history.",
])
add_feature_group(doc, "Tracking and communication", [
    "Shipment tracking by reference number with a clear status timeline.",
    "Estimated arrival and progress updates from booking through delivery.",
    "Push, email or WhatsApp-ready notification integration for key status changes.",
    "Customer support contact options and frequently asked questions.",
])
add_feature_group(doc, "Payments, documents and service information", [
    "Online payment integration readiness and payment confirmation records.",
    "Access to invoices, receipts and shipment summaries.",
    "Collection schedules, announcements and service availability notices.",
    "Customer feedback and service review submission.",
])

note = doc.add_table(rows=1, cols=1)
set_table_geometry(note, [9360])
set_cell_shading(note.cell(0, 0), LIGHT)
p = note.cell(0, 0).paragraphs[0]
p.paragraph_format.space_after = Pt(0)
set_run(p.add_run("Release package: "), size=10, color=BLUE, bold=True)
set_run(p.add_run("branded app icon and splash screen, production Android build, store screenshots and listing support, plus an iOS-ready build configuration where required."), size=10)

# PAGE 3 - staff app and shared foundation.
doc.add_page_break()
add_heading(doc, "2. Staff operations app", 1)
add_body(doc, "A role-based internal application for managers, administrators, finance personnel and drivers.")

add_feature_group(doc, "Management dashboard", [
    "At-a-glance shipment totals, active jobs, delivered shipments and revenue indicators.",
    "UK, Zimbabwe and combined operational filters.",
    "Attention alerts for pending collections, quotation requests and operational exceptions.",
])
add_feature_group(doc, "Shipment and customer operations", [
    "Searchable shipment list with status, customer, tracking and route filters.",
    "Shipment detail view with sender, receiver, package and journey information.",
    "Status updates with customer-notification triggers and audit history.",
    "Customer directory, manual booking and customer contact actions.",
])
add_feature_group(doc, "Finance and administration", [
    "Payments, invoices, receipts and finance dashboard access.",
    "Custom quotation review, response and status management.",
    "Collection schedule, delivery workflow, feedback and announcement management.",
    "Role-based menus for administrators, finance users and drivers.",
])
add_feature_group(doc, "Driver workflow", [
    "Daily collection and delivery run lists with customer and address details.",
    "Tap-to-call customer action and job status confirmation.",
    "Mark Collected and Mark Delivered workflows with secure database updates.",
])

add_heading(doc, "3. Shared platform and delivery", 1)
shared = doc.add_table(rows=1, cols=3)
set_table_geometry(shared, [3120, 3120, 3120])
for idx, (title, detail) in enumerate((
    ("Backend & data", "Secure cloud database, application APIs, file storage and synchronized customer/staff records."),
    ("Security", "Authenticated access, role permissions, validation, protected operational data and MFA readiness."),
    ("Quality & launch", "Responsive testing, core workflow testing, deployment configuration and release handover."),
)):
    cell = shared.cell(0, idx)
    set_cell_shading(cell, PALE_BLUE if idx != 1 else LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(title + "\n"), size=10.5, color=BLUE, bold=True)
    set_run(p.add_run(detail), size=9.4)

# PAGE 4 - commercial terms.
doc.add_page_break()
add_compact_heading(doc, "4. Project investment")
add_pricing_table(doc, [
    ("Customer mobile app - design, development and core workflows", "$450.00"),
    ("Staff operations app - management, finance and driver workflows", "$500.00"),
    ("Shared backend integration, roles, security and notifications", "$150.00"),
    ("Testing, production builds, store assets and release support", "$99.99"),
], "TOTAL FIXED PROJECT PRICE", "$1,199.99")

add_compact_heading(doc, "5. Recommended monthly service")
add_pricing_table(doc, [
    ("Cloud hosting and backend oversight", "$55.00"),
    ("Monitoring, backups and operational checks", "$30.00"),
    ("Maintenance and priority bug fixes", "$35.00"),
    ("Support and minor content/configuration updates", "$29.99"),
], "MONTHLY SERVICE TOTAL", "$149.99 / month")

add_body(doc, "The monthly service begins after production launch. It is within the requested $120-$180 monthly operating range and may be reviewed if usage, storage, messaging volume or support requirements increase materially.", italic=True, after=2)

add_compact_heading(doc, "6. Delivery, payment and assumptions")
add_terms_table(doc, [
    ("Estimated delivery", "6-8 weeks from approval, deposit and receipt of final content/access credentials."),
    ("Payment milestones", "$600.00 deposit to commence; $360.00 at functional beta; $239.99 before production release."),
    ("Included revisions", "Two reasonable review rounds during the agreed design and beta stages."),
    ("Warranty", "30-day post-launch correction period for defects in the agreed scope."),
    ("Quotation validity", "30 days from 16 July 2026."),
])

add_compact_heading(doc, "7. Exclusions and change control")
add_body(doc, "Unless expressly included above, the quotation excludes Google Play or Apple developer-account charges, payment-processor fees, SMS/WhatsApp/email usage charges, paid third-party services, hardware, data migration beyond supplied formats, and major features requested after scope approval. Additional work will be estimated and approved in writing before implementation.", after=2)

add_compact_heading(doc, "Acceptance")
add_body(doc, "Approval of this quotation confirms the scope, fixed project price and payment schedule described above. Final implementation will also be governed by a short project agreement and agreed delivery checklist.", after=2)

sig = doc.add_table(rows=1, cols=2)
set_table_geometry(sig, [4680, 4680])
for i, label in enumerate(("FOR ZIMBABWE SHIPPING", "FOR VOYAGE TECH")):
    p = sig.cell(0, i).paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(label + "\n"), size=8.5, color=MUTED, bold=True)
    set_run(p.add_run("Name / signature: ____________________\nDate: ______________"), size=8.8)
    set_cell_border(sig.cell(0, i), top={"val": "single", "sz": "5", "color": BORDER})

doc.core_properties.title = "Zimbabwe Shipping Mobile Apps - Commercial Quotation"
doc.core_properties.subject = "Quotation for customer and staff mobile applications"
doc.core_properties.author = "Voyage Tech"
doc.core_properties.company = "Voyage Tech"
doc.core_properties.comments = "Prepared for Zimbabwe Shipping"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
